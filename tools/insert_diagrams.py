# -*- coding: utf-8 -*-
import docx, struct
from docx.shared import Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH

DOC = r'D:\NienLuan\BaoCao_DoAn_HoanChinh.docx'
DIAGRAMS = r'D:\NienLuan\tools\diagrams'

mapping = [
    ('Sơ đồ kiến trúc tổng thể', 'architecture.png'),
    ('Sơ đồ ERD cơ sở dữ liệu', 'erd.png'),
    ('Sơ đồ Use Case tổng quan', 'usecase.png'),
    ('Activity Diagram luồng đặt vé', 'activity_booking.png'),
    ('Sequence Diagram', 'sequence_booking.png'),
]


def png_size(path):
    with open(path, 'rb') as f:
        head = f.read(26)
    return struct.unpack('>II', head[16:24])


def insert_image(par, path):
    w, h = png_size(path)
    wc = 15.0
    hc = wc * h / w
    if hc > 17.0:
        hc = 17.0
        wc = hc * w / h
    for r in list(par.runs):
        r._element.getparent().remove(r._element)
    par.alignment = WD_ALIGN_PARAGRAPH.CENTER
    par.add_run().add_picture(path, width=Cm(wc))


d = docx.Document(DOC)
cnt = 0
for p in list(d.paragraphs):
    t = p.text.strip()
    if not t.startswith('[CH'):
        continue
    for key, fname in mapping:
        if key in t:
            insert_image(p, DIAGRAMS + '\\' + fname)
            cnt += 1
            break

d.save(DOC)
print('inserted:', cnt)

d2 = docx.Document(DOC)
print('inline_shapes:', len(d2.inline_shapes))
ph = sum(1 for x in d2.paragraphs if x.text.strip().startswith('[CH'))
print('remaining_placeholders:', ph)
