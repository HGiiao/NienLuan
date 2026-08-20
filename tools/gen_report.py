# -*- coding: utf-8 -*-
"""
Sinh báo cáo đồ án .docx theo đúng quy định định dạng.
Cách chạy: python gen_report.py
"""
import docx
from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.enum.section import WD_SECTION
from docx.enum.style import WD_STYLE_TYPE
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

OUT = r"D:\NienLuan\BaoCao_DoAn_HoanChinh.docx"

FONT = "Times New Roman"

# ============================ helpers ============================
def set_font(run, size=13, bold=False, italic=False):
    run.font.name = FONT
    run.font.size = Pt(size)
    run.font.bold = bold
    run.font.italic = italic
    rPr = run._element.get_or_add_rPr()
    rFonts = rPr.find(qn('w:rFonts'))
    if rFonts is None:
        rFonts = OxmlElement('w:rFonts')
        rPr.append(rFonts)
    rFonts.set(qn('w:eastAsia'), FONT)

def para(doc, text="", size=13, bold=False, italic=False, align=WD_ALIGN_PARAGRAPH.JUSTIFY,
         space_after=6, line_multiple=1.3, keep_with_next=False):
    p = doc.add_paragraph()
    p.alignment = align
    pf = p.paragraph_format
    pf.line_spacing_rule = WD_LINE_SPACING.MULTIPLE
    pf.line_spacing = line_multiple
    pf.space_after = Pt(space_after)
    pf.space_before = Pt(0)
    if keep_with_next:
        pf.keep_with_next = True
    if text:
        r = p.add_run(text)
        set_font(r, size, bold, italic)
    return p

def chapter_title(doc, text):
    return para(doc, text, size=14, bold=True, align=WD_ALIGN_PARAGRAPH.LEFT,
                space_after=6, keep_with_next=True)

def heading(doc, text, level=1):
    # level 1 -> Heading1, level 2 -> Heading2, level 3 -> Heading3
    p = doc.add_paragraph(style=f"Heading {level}")
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    r = p.add_run(text)
    sizes = {1: 14, 2: 14, 3: 13}
    set_font(r, sizes.get(level, 13), bold=True)
    return p

def figure_placeholder(doc, label):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_after = Pt(2)
    r = p.add_run(label)
    set_font(r, 13, bold=False, italic=True)

def figure_caption(doc, text):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_after = Pt(6)
    r = p.add_run(text)
    set_font(r, 12, bold=False)

def table_caption(doc, text):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    p.paragraph_format.space_after = Pt(2)
    p.paragraph_format.keep_with_next = True
    r = p.add_run(text)
    set_font(r, 13, bold=True)

def make_table(doc, headers, rows, widths=None, font_size=12):
    t = doc.add_table(rows=1, cols=len(headers))
    t.style = 'Table Grid'
    hdr = t.rows[0].cells
    for i, h in enumerate(headers):
        hdr[i].text = ""
        p = hdr[i].paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r = p.add_run(h)
        set_font(r, font_size, bold=True)
    for row in rows:
        cells = t.add_row().cells
        for i, val in enumerate(row):
            cells[i].text = ""
            p = cells[i].paragraphs[0]
            p.alignment = WD_ALIGN_PARAGRAPH.LEFT
            p.paragraph_format.space_after = Pt(0)
            r = p.add_run(str(val))
            set_font(r, font_size)
    if widths:
        for i, w in enumerate(widths):
            for row in t.rows:
                row.cells[i].width = Cm(w)
    # spacing after table
    sp = doc.add_paragraph()
    sp.paragraph_format.space_after = Pt(6)
    return t

def add_field(paragraph, field_code):
    r = paragraph.add_run()
    set_font(r, 12)
    fldChar1 = OxmlElement('w:fldChar'); fldChar1.set(qn('w:fldCharType'), 'begin')
    instrText = OxmlElement('w:instrText'); instrText.set(qn('xml:space'), 'preserve')
    instrText.text = field_code
    fldChar2 = OxmlElement('w:fldChar'); fldChar2.set(qn('w:fldCharType'), 'separate')
    t = OxmlElement('w:t'); t.text = ""
    r._element.append(fldChar1)
    r._element.append(instrText)
    r._element.append(fldChar2)
    r._element.append(t)
    # end
    r2 = paragraph.add_run()
    set_font(r2, 12)
    fldChar3 = OxmlElement('w:fldChar'); fldChar3.set(qn('w:fldCharType'), 'end')
    r2._element.append(fldChar3)

def page_break(doc):
    p = doc.add_paragraph()
    run = p.add_run()
    run.add_break(docx.enum.text.WD_BREAK.PAGE)

# ============================ document setup ============================
doc = Document()

# default style
normal = doc.styles['Normal']
normal.font.name = FONT
normal.font.size = Pt(13)
normal.element.rPr.rFonts.set(qn('w:eastAsia'), FONT)
normal.paragraph_format.line_spacing_rule = WD_LINE_SPACING.MULTIPLE
normal.paragraph_format.line_spacing = 1.3
normal.paragraph_format.space_after = Pt(6)

# heading styles font override
for lvl, sz in [(1,14),(2,14),(3,13)]:
    st = doc.styles[f'Heading {lvl}']
    st.font.name = FONT
    st.font.size = Pt(sz)
    st.font.bold = True
    st.font.color.rgb = RGBColor(0,0,0)
    st.paragraph_format.space_after = Pt(6)
    st.paragraph_format.space_before = Pt(6)
    st.paragraph_format.line_spacing_rule = WD_LINE_SPACING.MULTIPLE
    st.paragraph_format.line_spacing = 1.3
    rpr = st.element.get_or_add_rPr()
    rfonts = rpr.find(qn('w:rFonts'))
    if rfonts is None:
        rfonts = OxmlElement('w:rFonts'); rpr.append(rfonts)
    rfonts.set(qn('w:eastAsia'), FONT)

# margins + A4 for all sections (set later per section too)
def set_section(section):
    section.page_width = Cm(21.0)
    section.page_height = Cm(29.7)
    section.left_margin = Cm(3.0)
    section.right_margin = Cm(2.0)
    section.top_margin = Cm(2.0)
    section.bottom_margin = Cm(2.0)

first = doc.sections[0]
set_section(first)

# footer page number field helper
def enable_page_number(section):
    footer = section.footer
    footer.is_linked_to_previous = False
    p = footer.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.text = ""
    add_field(p, "PAGE")

def blank_footer(section):
    footer = section.footer
    footer.is_linked_to_previous = False
    p = footer.paragraphs[0]
    p.text = ""
    for r in list(p.runs):
        r.text = ""

# ============================ COVER (main) ============================
sec = doc.sections[0]
blank_footer(sec)
para(doc, "TRƯỜNG ĐẠI HỌC ĐỒNG THÁP", size=13, bold=True, align=WD_ALIGN_PARAGRAPH.CENTER, space_after=0)
para(doc, "KHOA CÔNG NGHỆ - KỸ THUẬT", size=13, bold=True, align=WD_ALIGN_PARAGRAPH.CENTER, space_after=0)
para(doc, "", size=13, align=WD_ALIGN_PARAGRAPH.CENTER, space_after=0)
para(doc, "", size=13, align=WD_ALIGN_PARAGRAPH.CENTER, space_after=0)
para(doc, "ĐỒ ÁN MÔN HỌC 2", size=18, bold=True, align=WD_ALIGN_PARAGRAPH.CENTER, space_after=0)
para(doc, "", size=13, align=WD_ALIGN_PARAGRAPH.CENTER, space_after=0)
para(doc, "", size=13, align=WD_ALIGN_PARAGRAPH.CENTER, space_after=0)
para(doc, "WEBSITE TỔNG HỢP DỮ LIỆU VÉ MÁY BAY,", size=14, bold=True, align=WD_ALIGN_PARAGRAPH.CENTER, space_after=0)
para(doc, "VÉ TÀU HỎA VÀ GỢI Ý SO SÁNH GIÁ VÉ", size=14, bold=True, align=WD_ALIGN_PARAGRAPH.CENTER, space_after=0)
para(doc, "THEO THỜI GIAN THỰC", size=14, bold=True, align=WD_ALIGN_PARAGRAPH.CENTER, space_after=0)
para(doc, "", size=13, align=WD_ALIGN_PARAGRAPH.CENTER, space_after=0)
para(doc, "", size=13, align=WD_ALIGN_PARAGRAPH.CENTER, space_after=0)
para(doc, "", size=13, align=WD_ALIGN_PARAGRAPH.CENTER, space_after=0)
para(doc, "Sinh viên thực hiện:", size=13, bold=True, align=WD_ALIGN_PARAGRAPH.CENTER, space_after=0)
para(doc, "NGUYỄN THỊ HUỲNH GIAO", size=13, bold=True, align=WD_ALIGN_PARAGRAPH.CENTER, space_after=0)
para(doc, "MSSV: 0023412018", size=13, align=WD_ALIGN_PARAGRAPH.CENTER, space_after=0)
para(doc, "Lớp: ĐHCNTT23B-IT", size=13, align=WD_ALIGN_PARAGRAPH.CENTER, space_after=0)
para(doc, "Ngành: Công Nghệ Thông Tin", size=13, align=WD_ALIGN_PARAGRAPH.CENTER, space_after=0)
para(doc, "", size=13, align=WD_ALIGN_PARAGRAPH.CENTER, space_after=0)
para(doc, "", size=13, align=WD_ALIGN_PARAGRAPH.CENTER, space_after=0)
para(doc, "Đồng Tháp, ngày 21 tháng 08 năm 2026", size=13, align=WD_ALIGN_PARAGRAPH.CENTER, space_after=0)
page_break(doc)

# ============================ COVER LINER (bìa lót) ============================
sec = doc.add_section(WD_SECTION.NEW_PAGE)
set_section(sec)
blank_footer(sec)
para(doc, "TRƯỜNG ĐẠI HỌC ĐỒNG THÁP", size=13, bold=True, align=WD_ALIGN_PARAGRAPH.CENTER, space_after=0)
para(doc, "KHOA CÔNG NGHỆ - KỸ THUẬT", size=13, bold=True, align=WD_ALIGN_PARAGRAPH.CENTER, space_after=0)
para(doc, "", size=13, align=WD_ALIGN_PARAGRAPH.CENTER, space_after=0)
para(doc, "", size=13, align=WD_ALIGN_PARAGRAPH.CENTER, space_after=0)
para(doc, "ĐỒ ÁN MÔN HỌC 2", size=16, bold=True, align=WD_ALIGN_PARAGRAPH.CENTER, space_after=0)
para(doc, "", size=13, align=WD_ALIGN_PARAGRAPH.CENTER, space_after=0)
para(doc, "WEBSITE TỔNG HỢP DỮ LIỆU VÉ MÁY BAY, VÉ TÀU HỎA", size=14, bold=True, align=WD_ALIGN_PARAGRAPH.CENTER, space_after=0)
para(doc, "VÀ GỢI Ý SO SÁNH GIÁ VÉ THEO THỜI GIAN THỰC", size=14, bold=True, align=WD_ALIGN_PARAGRAPH.CENTER, space_after=0)
para(doc, "", size=13, align=WD_ALIGN_PARAGRAPH.CENTER, space_after=0)
para(doc, "Mã số đề tài: 73", size=13, bold=True, align=WD_ALIGN_PARAGRAPH.CENTER, space_after=0)
para(doc, "Ngành: Công Nghệ Thông Tin", size=13, align=WD_ALIGN_PARAGRAPH.CENTER, space_after=0)
para(doc, "Lớp: ĐHCNTT23B-IT", size=13, align=WD_ALIGN_PARAGRAPH.CENTER, space_after=0)
para(doc, "", size=13, align=WD_ALIGN_PARAGRAPH.CENTER, space_after=0)
para(doc, "", size=13, align=WD_ALIGN_PARAGRAPH.CENTER, space_after=0)
para(doc, "Sinh viên thực hiện:", size=13, bold=True, align=WD_ALIGN_PARAGRAPH.CENTER, space_after=0)
para(doc, "Nguyễn Thị Huỳnh Giao", size=13, align=WD_ALIGN_PARAGRAPH.CENTER, space_after=0)
para(doc, "MSSV: 0023412018", size=13, align=WD_ALIGN_PARAGRAPH.CENTER, space_after=0)
para(doc, "", size=13, align=WD_ALIGN_PARAGRAPH.CENTER, space_after=0)
para(doc, "Giảng viên hướng dẫn: [CẦN BỔ SUNG]", size=13, align=WD_ALIGN_PARAGRAPH.CENTER, space_after=0)
para(doc, "", size=13, align=WD_ALIGN_PARAGRAPH.CENTER, space_after=0)
para(doc, "", size=13, align=WD_ALIGN_PARAGRAPH.CENTER, space_after=0)
para(doc, "Đồng Tháp, năm 2026", size=13, align=WD_ALIGN_PARAGRAPH.CENTER, space_after=0)
page_break(doc)

# ============================ LỜI CẢM ƠN ============================
sec = doc.add_section(WD_SECTION.NEW_PAGE)
set_section(sec)
blank_footer(sec)
chapter_title(doc, "LỜI CẢM ƠN")
para(doc, "Trước tiên, em xin bày tỏ lòng biết ơn sâu sắc đến quý thầy cô Trường Đại học Đồng Tháp, đặc biệt là các thầy cô của Khoa Công nghệ - Kỹ thuật, đã tận tình giảng dạy và truyền đạt những kiến thức quý báu trong suốt thời gian em học tập tại trường. Những kiến thức về lập trình web, cơ sở dữ liệu, phân tích thiết kế hệ thống và công nghệ phần mềm là nền tảng quan trọng để em hoàn thành đồ án môn học này.")
para(doc, "Em xin chân thành cảm ơn giảng viên hướng dẫn [CẦN BỔ SUNG] đã dành thời gian định hướng, góp ý và giúp em hoàn thiện đề tài. Sự hướng dẫn tận tình của thầy/cô là động lực lớn giúp em vượt qua những khó khăn trong quá trình thực hiện.")
para(doc, "Em cũng xin cảm ơn gia đình và bạn bè đã luôn đồng hành, động viên và hỗ trợ em trong suốt quá trình học tập và thực hiện đồ án.")
para(doc, "Do thời gian và kiến thức còn hạn chế, đồ án khó tránh khỏi những thiếu sót. Em rất mong nhận được sự góp ý của quý thầy cô để đồ án được hoàn thiện hơn.")
para(doc, "Em xin trân trọng cảm ơn!")
page_break(doc)

# ============================ LỜI DẪN ============================
sec = doc.add_section(WD_SECTION.NEW_PAGE)
set_section(sec)
blank_footer(sec)
chapter_title(doc, "LỜI DẪN")
para(doc, "Trong bối cảnh nhu cầu đi lại của người dân ngày càng tăng, việc so sánh và lựa chọn phương tiện di chuyển với chi phí hợp lý trở thành nhu cầu thiết yếu. Người dùng thường phải truy cập nhiều website khác nhau để tra cứu giá vé máy bay, vé tàu hỏa và xe khách, dẫn đến tốn thời gian và khó so sánh. Đề tài \u201cWebsite tổng hợp dữ liệu vé máy bay, vé tàu hỏa và gợi ý so sánh giá vé theo thời gian thực\u201d được thực hiện nhằm giải quyết vấn đề này.")
para(doc, "Hệ thống Vé247 được xây dựng với mục đích tổng hợp dữ liệu giá vé của ba loại phương tiện gồm máy bay, tàu hỏa và xe khách vào một nền tảng duy nhất, cho phép người dùng tra cứu, so sánh giá, xem biểu đồ xu hướng giá và nhận gợi ý lộ trình kết hợp tối ưu chi phí. Bên cạnh đó, hệ thống còn hỗ trợ đặt vé trực tiếp hoặc chuyển hướng người dùng đến trang đặt vé chính thức của hãng.")
para(doc, "Công nghệ chính sử dụng trong đồ án gồm: frontend dùng React, Vite, Tailwind CSS, recharts và SignalR client; backend dùng ASP.NET Core (.NET 10), Entity Framework Core, SignalR; dữ liệu được lưu trữ trên SQL Server (Azure SQL Server). Các chức năng chính gồm tra cứu vé, so sánh giá theo thời gian thực (mô phỏng), gợi ý lộ trình tối ưu, đặt vé, thanh toán đa cổng, cảnh báo giá, quản trị hệ thống và nhiều tính năng hỗ trợ khác.")
para(doc, "Kết quả đạt được là một website hoàn chỉnh, vận hành được với đầy đủ các chức năng cốt lõi của một nền tảng tổng hợp vé, đáp ứng các mục tiêu đề ra của đồ án.")
page_break(doc)

# ============================ MỤC LỤC ============================
sec = doc.add_section(WD_SECTION.NEW_PAGE)
set_section(sec)
blank_footer(sec)
chapter_title(doc, "MỤC LỤC")
p = doc.add_paragraph()
add_field(p, "TOC \\o \"1-3\" \\h \\z \\u")
page_break(doc)

# ============================ DANH MỤC HÌNH ============================
sec = doc.add_section(WD_SECTION.NEW_PAGE)
set_section(sec)
blank_footer(sec)
chapter_title(doc, "DANH MỤC HÌNH")
p = doc.add_paragraph()
add_field(p, "TOC \\h \\z \\c \"Hình\"")
page_break(doc)

# ============================ DANH MỤC BẢNG ============================
sec = doc.add_section(WD_SECTION.NEW_PAGE)
set_section(sec)
blank_footer(sec)
chapter_title(doc, "DANH MỤC BẢNG")
p = doc.add_paragraph()
add_field(p, "TOC \\h \\z \\c \"Bảng\"")
page_break(doc)

# ============================ PHẦN MỘT: MỞ ĐẦU ============================
sec = doc.add_section(WD_SECTION.NEW_PAGE)
set_section(sec)
# bắt đầu đánh số trang từ phần này, reset về 1
sec.footer.is_linked_to_previous = False
# reset page numbering to 1
sectPr = sec._sectPr
pgNumType = sectPr.find(qn('w:pgNumType'))
if pgNumType is None:
    pgNumType = OxmlElement('w:pgNumType'); sectPr.append(pgNumType)
pgNumType.set(qn('w:start'), '1')
enable_page_number(sec)

chapter_title(doc, "PHẦN MỘT: MỞ ĐẦU")

heading(doc, "1. Đặt vấn đề", 1)
para(doc, "Hiện nay, nhu cầu di chuyển của người dân ngày càng đa dạng với nhiều phương tiện khác nhau như máy bay, tàu hỏa và xe khách. Tuy nhiên, mỗi hãng vận tải thường chỉ cung cấp dịch vụ tra cứu và đặt vé trên website riêng của mình. Điều này khiến người dùng gặp khó khăn khi muốn so sánh giá vé giữa các phương tiện, các hãng khác nhau để lựa chọn phương án tối ưu về chi phí và thời gian.")
para(doc, "Bên cạnh đó, giá vé máy bay thường biến động mạnh theo thời gian và theo mùa. Người dùng thiếu thông tin về xu hướng giá sẽ khó đưa ra quyết định đặt vé hợp lý. Việc phải mở nhiều tab trình duyệt, nhập đi nhập lại thông tin điểm đi, điểm đến, ngày đi ở từng website là bất tiện và tốn thời gian.")
para(doc, "Xuất phát từ thực trạng trên, đề tài xây dựng một website tổng hợp dữ liệu giá vé từ nhiều nguồn, tập trung vào ba loại phương tiện chính, giúp người dùng tra cứu nhanh, so sánh giá, theo dõi xu hướng giá và nhận gợi ý lộ trình tối ưu, đồng thời hỗ trợ đặt vé hoặc chuyển hướng đến trang đặt vé chính thức.")

heading(doc, "2. Ý nghĩa lý thuyết và thực tiễn của đề tài", 1)
heading(doc, "2.1. Ý nghĩa lý thuyết", 2)
para(doc, "Đề tài vận dụng các kiến thức về phát triển web full-stack, bao gồm xây dựng giao diện với React, xây dựng API với ASP.NET Core, thao tác cơ sở dữ liệu với Entity Framework Core, và giao tiếp thời gian thực với SignalR. Đồng thời, đề tài áp dụng các kỹ thuật phân tích, thiết kế hệ thống, lập trình hướng đối tượng và tổ chức dữ liệu, là cơ sở để sinh viên củng cố và mở rộng kiến thức đã học.")
heading(doc, "2.2. Ý nghĩa thực tiễn", 2)
para(doc, "Hệ thống giúp người dùng tiết kiệm thời gian tra cứu giá vé, dễ dàng so sánh giữa các phương tiện và hãng vận tải, nhận gợi ý lộ trình tối ưu chi phí, và có thể đặt vé trực tiếp hoặc được dẫn đến trang đặt vé chính thức. Từ đó nâng cao trải nghiệm lựa chọn phương tiện di chuyển của người dùng.")

heading(doc, "3. Mục tiêu của đề tài", 1)
heading(doc, "3.1. Mục tiêu tổng quát", 2)
para(doc, "Xây dựng một website tổng hợp dữ liệu giá vé máy bay, vé tàu hỏa và xe khách, cho phép tra cứu, so sánh giá theo thời gian thực (mô phỏng) và gợi ý lộ trình kết hợp tối ưu chi phí.")
heading(doc, "3.2. Mục tiêu cụ thể", 2)
para(doc, "Các mục tiêu cụ thể của đề tài gồm: (1) xây dựng chức năng tra cứu vé máy bay, tàu hỏa, xe khách với bộ lọc và sắp xếp; (2) xây dựng chức năng so sánh giá và hiển thị biểu đồ xu hướng giá; (3) xây dựng chức năng gợi ý lộ trình kết hợp tối ưu chi phí; (4) xây dựng chức năng đặt vé trực tiếp và chuyển hướng đến trang đặt vé chính thức; (5) xây dựng các chức năng hỗ trợ như thanh toán, cảnh báo giá, quản trị hệ thống.")

heading(doc, "4. Phạm vi của đề tài", 1)
heading(doc, "4.1. Phạm vi chức năng", 2)
para(doc, "Hệ thống tập trung vào các chức năng: tra cứu và so sánh giá vé, gợi ý lộ trình tối ưu, đặt vé và thanh toán, cảnh báo giá, quản trị hệ thống, cùng các tính năng hỗ trợ khác. Dữ liệu giá vé hiện tại là dữ liệu mẫu được khởi tạo vào cơ sở dữ liệu.")
heading(doc, "4.2. Phạm vi người dùng", 2)
para(doc, "Hệ thống phục vụ hai nhóm người dùng chính: khách truy cập/người dùng đã đăng ký có nhu cầu tra cứu và đặt vé, và quản trị viên có nhiệm vụ quản lý dữ liệu và vận hành hệ thống.")
heading(doc, "4.3. Phạm vi công nghệ", 2)
para(doc, "Frontend sử dụng React, Vite, Tailwind CSS, React Router, recharts, SignalR client, Clerk. Backend sử dụng ASP.NET Core (.NET 10), Entity Framework Core, SignalR, BCrypt. Dữ liệu lưu trên SQL Server (Azure SQL Server). Kiểm thử đơn vị dùng xUnit.")
heading(doc, "4.4. Phạm vi triển khai", 2)
para(doc, "Hệ thống được vận hành và kiểm thử trong môi trường phát triển cục bộ (backend chạy trên cổng 5000, frontend chạy trên cổng 5173). Hệ thống được định hướng triển khai trên nền tảng Vercel. Tại thời điểm thực hiện báo cáo, quá trình triển khai thực tế trên Vercel chưa được hoàn tất.")

heading(doc, "5. Các phương pháp nghiên cứu", 1)
para(doc, "Đề tài sử dụng các phương pháp: nghiên cứu tài liệu về công nghệ web; phân tích yêu cầu từ bài toán thực tế; phân tích và thiết kế hệ thống; nghiên cứu và lựa chọn công nghệ; thiết kế giao diện; lập trình và tích hợp; kiểm thử chức năng; và đánh giá kết quả đạt được.")

heading(doc, "6. Kế hoạch thực hiện đề tài", 1)
table_caption(doc, "Bảng 0.1. Kế hoạch thực hiện đề tài")
make_table(doc,
    ["STT", "Giai đoạn", "Nội dung", "Kết quả"],
    [["1", "Khảo sát", "Phân tích yêu cầu và bài toán thực tế", "Xác định yêu cầu"],
     ["2", "Phân tích", "Phân tích hệ thống và dữ liệu", "Đặc tả hệ thống"],
     ["3", "Thiết kế", "Thiết kế giao diện và kiến trúc", "Thiết kế hệ thống"],
     ["4", "Lập trình", "Xây dựng các chức năng frontend/backend", "Source code"],
     ["5", "Kiểm thử", "Kiểm thử chức năng và tích hợp", "Kết quả kiểm thử"],
     ["6", "Hoàn thiện", "Hoàn thiện và viết báo cáo", "Báo cáo đồ án"]],
    widths=[1.5, 3.0, 6.5, 4.0])
page_break(doc)

# ============================ PHẦN HAI: NỘI DUNG ============================
sec = doc.add_section(WD_SECTION.NEW_PAGE)
set_section(sec)
sec.footer.is_linked_to_previous = False
enable_page_number(sec)

chapter_title(doc, "PHẦN HAI: NỘI DUNG")

# ---------------- CHƯƠNG 1 ----------------
chapter_title(doc, "CHƯƠNG 1: CƠ SỞ LÝ THUYẾT")

heading(doc, "1.1. Các công trình/hệ thống liên quan", 1)
para(doc, "Trên thị trường hiện có nhiều nền tảng so sánh và đặt vé trực tuyến như các website bán vé của hãng hàng không, website của Đường sắt Việt Nam, hay các đại lý vé trực tuyến. Tuy nhiên, hầu hết các nền tảng này chỉ tập trung vào một loại phương tiện hoặc một hãng vận tải. Một số nền tảng tổng hợp quốc tế cung cấp dịch vụ so sánh giá vé máy bay, nhưng chưa phổ biến với thị trường nội địa Việt Nam và thường yêu cầu chi phí sử dụng dịch vụ API cao. Vì vậy, đề tài hướng đến xây dựng một nền tảng tổng hợp ba loại phương tiện trong nước, phù hợp với nhu cầu người dùng Việt Nam.")

heading(doc, "1.2. Tổng quan về các công nghệ sử dụng", 1)
para(doc, "Dựa trên source code thực tế, hệ thống sử dụng các công nghệ chính sau đây.")

heading(doc, "1.2.1. React", 2)
para(doc, "React là thư viện JavaScript phổ biến để xây dựng giao diện người dùng dựa trên thành phần (component). Trong đề tài, React được dùng để xây dựng toàn bộ giao diện frontend với các component, page và quản lý state.")
heading(doc, "1.2.2. Vite", 2)
para(doc, "Vite là công cụ build cho ứng dụng web hiện đại, cung cấp máy chủ phát triển nhanh và khả năng đóng gói tối ưu. Trong đề tài, Vite được dùng để chạy và build frontend.")
heading(doc, "1.2.3. Tailwind CSS", 2)
para(doc, "Tailwind CSS là framework CSS theo kiểu utility-first giúp xây dựng giao diện nhanh và nhất quán. Đề tài dùng Tailwind CSS để thiết kế giao diện.")
heading(doc, "1.2.4. ASP.NET Core (.NET 10)", 2)
para(doc, "ASP.NET Core là framework mã nguồn mở của Microsoft để xây dựng ứng dụng web và API. Trong đề tài, ASP.NET Core được dùng để xây dựng REST API backend.")
heading(doc, "1.2.5. Entity Framework Core", 2)
para(doc, "Entity Framework Core là ORM (Object-Relational Mapping) của .NET, giúp thao tác cơ sở dữ liệu bằng đối tượng C#. Đề tài dùng EF Core kết nối và truy vấn SQL Server.")
heading(doc, "1.2.6. SQL Server / Azure SQL Server", 2)
para(doc, "SQL Server là hệ quản trị cơ sở dữ liệu quan hệ của Microsoft. Đề tài lưu trữ dữ liệu trên SQL Server (Azure SQL Server) với các bảng chuyến bay, tàu hỏa, xe khách, đặt chỗ, người dùng...")
heading(doc, "1.2.7. SignalR", 2)
para(doc, "SignalR là thư viện của ASP.NET Core hỗ trợ giao tiếp thời gian thực qua WebSocket. Đề tài dùng SignalR để đẩy cập nhật giá vé theo thời gian thực (mô phỏng) đến trình duyệt.")
heading(doc, "1.2.8. recharts", 2)
para(doc, "recharts là thư viện biểu đồ cho React. Đề tài dùng recharts để vẽ biểu đồ xu hướng giá vé.")
heading(doc, "1.2.9. Các thư viện hỗ trợ khác", 2)
para(doc, "Bên cạnh đó, đề tài sử dụng React Router (định tuyến), framer-motion (hiệu ứng), Clerk (xác thực), axios (gọi API), BCrypt (băm mật khẩu), xUnit (kiểm thử đơn vị) và các thư viện hỗ trợ thanh toán VNPay, PayOS, MoMo, ZaloPay, VietQR.")

heading(doc, "1.3. Kiến trúc hệ thống", 1)
para(doc, "Hệ thống được xây dựng theo kiến trúc ba tầng gồm Frontend, Backend và Database. Frontend (React) giao tiếp với Backend (ASP.NET Core REST API) thông qua giao thức HTTP, và nhận cập nhật thời gian thực qua SignalR. Backend thao tác dữ liệu với cơ sở dữ liệu SQL Server thông qua Entity Framework Core.")
figure_placeholder(doc, "[CHÈN ẢNH: Sơ đồ kiến trúc tổng thể]")
figure_caption(doc, "Hình 1.1. Kiến trúc tổng thể của hệ thống")

heading(doc, "1.4. Cơ sở dữ liệu", 1)
para(doc, "Hệ thống sử dụng SQL Server (Azure SQL Server) với các bảng chính sau: Flights (chuyến bay), Trains (tàu hỏa), Buses (xe khách), Users (người dùng), Bookings (đặt chỗ), BookingSegments (chặng của đặt chỗ), PriceHistories (lịch sử giá), PriceAlerts (cảnh báo giá), PromoCodes (mã giảm giá), Reviews (đánh giá), Notifications (thông báo), và nhiều bảng hỗ trợ khác như SubscriptionPlans, UserSubscriptions, InsurancePackages, LuckyWheelSpins, CommunityTips...")
para(doc, "Các bảng dữ liệu chính có khóa chính và các quan hệ khóa ngoại. Ví dụ, bảng Bookings tham chiếu đến Users, Flights, Trains, Buses; bảng BookingSegments tham chiếu đến Bookings. Các bảng chuyến bay có chỉ mục theo tuyến đường và ngày để tăng tốc truy vấn.")
figure_placeholder(doc, "[CHÈN ẢNH: Sơ đồ ERD cơ sở dữ liệu]")
figure_caption(doc, "Hình 1.2. Sơ đồ thực thể - quan hệ cơ sở dữ liệu")
page_break(doc)

# ---------------- CHƯƠNG 2 ----------------
chapter_title(doc, "CHƯƠNG 2: PHÂN TÍCH VÀ THIẾT KẾ HỆ THỐNG")

heading(doc, "2.1. Đặc tả bài toán", 1)
para(doc, "Bài toán đặt ra là xây dựng một website tổng hợp dữ liệu giá vé của ba loại phương tiện (máy bay, tàu hỏa, xe khách), cho phép người dùng tra cứu vé theo điểm đi, điểm đến và ngày đi, so sánh giá giữa các phương tiện, xem xu hướng giá, nhận gợi ý lộ trình kết hợp tối ưu chi phí, và thực hiện đặt vé trực tiếp hoặc được chuyển hướng đến trang đặt vé chính thức. Hệ thống cũng cần có trang quản trị để quản lý dữ liệu và vận hành.")

heading(doc, "2.2. Đối tượng sử dụng hệ thống", 1)
table_caption(doc, "Bảng 2.1. Các đối tượng sử dụng hệ thống")
make_table(doc,
    ["Đối tượng", "Mô tả"],
    [["Khách / Người dùng", "Tra cứu, so sánh giá, đặt vé, thanh toán, cảnh báo giá"],
     ["Quản trị viên", "Quản lý chuyến bay, tàu, xe, người dùng, đặt chỗ, mã giảm giá, thống kê, phát thông báo"]],
    widths=[4.5, 10.5])

heading(doc, "2.3. Yêu cầu chức năng", 1)
table_caption(doc, "Bảng 2.2. Danh sách yêu cầu chức năng chính")
make_table(doc,
    ["STT", "Mã chức năng", "Tên chức năng", "Mô tả"],
    [["1", "F01", "Tra cứu chuyến bay", "Tìm kiếm chuyến bay theo tuyến, ngày với bộ lọc và sắp xếp"],
     ["2", "F02", "Tra cứu tàu hỏa", "Tìm kiếm chuyến tàu theo tuyến, ngày với bộ lọc và sắp xếp"],
     ["3", "F03", "Tra cứu xe khách", "Tìm kiếm chuyến xe khách theo tuyến, ngày"],
     ["4", "F04", "So sánh giá", "So sánh giá vé giữa máy bay, tàu, xe trên cùng tuyến"],
     ["5", "F05", "Biểu đồ xu hướng giá", "Hiển thị biểu đồ xu hướng giá và dự đoán giá"],
     ["6", "F06", "Gợi ý lộ trình tối ưu", "Gợi ý lộ trình kết hợp nhiều phương tiện tối ưu chi phí"],
     ["7", "F07", "Đặt vé", "Tạo đặt chỗ cho chuyến bay/tàu/xe"],
     ["8", "F08", "Thanh toán", "Thanh toán đặt chỗ qua các cổng VNPay, PayOS, MoMo, ZaloPay, VietQR"],
     ["9", "F09", "Chuyển hướng đặt vé chính thức", "Chuyển người dùng đến trang đặt vé chính thức của hãng"],
     ["10", "F10", "Xác thực người dùng", "Đăng ký, đăng nhập, xác minh email"],
     ["11", "F11", "Cảnh báo giá", "Đặt và nhận cảnh báo khi giá giảm"],
     ["12", "F12", "Quản trị hệ thống", "Quản lý dữ liệu, người dùng, thống kê, thông báo"]],
    widths=[1.2, 2.8, 5.0, 6.0])

heading(doc, "2.4. Yêu cầu phi chức năng", 1)
para(doc, "Hệ thống hướng đến các yêu cầu phi chức năng: dễ sử dụng với giao diện thân thiện và responsive; bảo mật với mã hóa mật khẩu bằng BCrypt và kiểm soát quyền truy cập quản trị; hiệu năng truy vấn được tối ưu bằng chỉ mục trên các bảng; khả năng mở rộng nhờ kiến trúc tách biệt frontend/backend; và khả năng triển khai lên nền tảng đám mây (định hướng Vercel cho frontend).")

heading(doc, "2.5. Use Case", 1)
figure_placeholder(doc, "[CHÈN ẢNH: Sơ đồ Use Case tổng quan]")
figure_caption(doc, "Hình 2.1. Sơ đồ Use Case tổng quan")
table_caption(doc, "Bảng 2.3. Mô tả các Use Case chính")
make_table(doc,
    ["Use Case", "Actor", "Mô tả", "Tiền điều kiện", "Kết quả"],
    [["Đăng nhập", "Người dùng", "Xác thực tài khoản", "Có tài khoản", "Vào hệ thống"],
     ["Tìm kiếm vé", "Người dùng", "Tìm chuyến bay/tàu/xe", "Chọn tuyến, ngày", "Danh sách vé"],
     ["So sánh giá", "Người dùng", "So sánh giá các phương tiện", "Chọn tuyến, ngày", "Kết quả so sánh"],
     ["Đặt vé", "Người dùng", "Tạo đặt chỗ", "Đã chọn vé", "Đặt chỗ thành công"],
     ["Thanh toán", "Người dùng", "Thanh toán đặt chỗ", "Có đặt chỗ", "Xác nhận thanh toán"],
     ["Quản lý dữ liệu", "Quản trị viên", "Quản lý chuyến bay/tàu/xe", "Đăng nhập quản trị", "Cập nhật dữ liệu"]],
    widths=[3.2, 2.5, 4.0, 3.0, 3.0])

heading(doc, "2.6. Activity Diagram", 1)
para(doc, "Các nghiệp vụ quan trọng của hệ thống bao gồm luồng đăng nhập, tìm kiếm vé, đặt vé, thanh toán và quản lý dữ liệu. Trong đó, luồng đặt vé và thanh toán là nghiệp vụ phức tạp nhất.")
figure_placeholder(doc, "[CHÈN ẢNH: Activity Diagram luồng đặt vé và thanh toán]")
figure_caption(doc, "Hình 2.2. Activity Diagram luồng đặt vé")

heading(doc, "2.7. Sequence Diagram", 1)
para(doc, "Sequence Diagram minh họa luồng tương tác giữa người dùng, frontend, backend và cơ sở dữ liệu cho chức năng đặt vé.")
figure_placeholder(doc, "[CHÈN ẢNH: Sequence Diagram đặt vé]")
figure_caption(doc, "Hình 2.3. Sequence Diagram chức năng đặt vé")

heading(doc, "2.8. Thiết kế dữ liệu", 1)
para(doc, "Các entity chính trong hệ thống gồm User, Flight, Train, Bus, Booking, BookingSegment, PriceHistory, PriceAlert, PromoCode, Review, Notification và nhiều entity hỗ trợ khác. Dữ liệu được tổ chức thành các bảng có khóa chính, khóa ngoại và chỉ mục phù hợp.")
para(doc, "Ví dụ, entity Booking chứa thông tin người dùng, chuyến bay/tàu/xe, tổng giá, trạng thái, phương thức thanh toán, mã giao dịch; entity BookingSegment mô tả từng chặng trong một lộ trình kết hợp nhiều phương tiện.")
page_break(doc)

# ---------------- CHƯƠNG 3 ----------------
chapter_title(doc, "CHƯƠNG 3: XÂY DỰNG CHƯƠNG TRÌNH")

heading(doc, "3.1. Môi trường phát triển", 1)
table_caption(doc, "Bảng 3.1. Môi trường phát triển")
make_table(doc,
    ["Thành phần", "Công nghệ/Phiên bản", "Vai trò"],
    [["Frontend", "React 19, Vite 8, Tailwind CSS 4", "Giao diện người dùng"],
     ["Backend", "ASP.NET Core (.NET 10)", "REST API"],
     ["ORM", "Entity Framework Core 10", "Truy vấn dữ liệu"],
     ["Database", "SQL Server / Azure SQL Server", "Lưu trữ dữ liệu"],
     ["Thời gian thực", "SignalR", "Cập nhật giá thời gian thực"],
     ["Xác thực", "Clerk, BCrypt", "Đăng nhập, băm mật khẩu"],
     ["Kiểm thử", "xUnit", "Kiểm thử đơn vị"]],
    widths=[4.0, 6.5, 4.5])

heading(doc, "3.2. Cấu trúc project", 1)
para(doc, "Project gồm hai phần chính là frontend và backend. Frontend nằm trong thư mục frontend, backend nằm trong thư mục backend.")
para(doc, "Thư mục frontend/src chứa: components (các thành phần giao diện), pages (các trang), services (gọi API), hooks (custom hook), ui (thành phần giao diện dùng chung), admin (giao diện quản trị), config, utils.")
para(doc, "Thư mục backend/FlightAggregatorApi chứa: Controllers (các controller API), Models (các entity), Services (nghiệp vụ), Data (DbContext), Hubs (SignalR), Helpers. Thư mục backend/FlightAggregatorApi.Tests chứa các bài kiểm thử đơn vị.")

heading(doc, "3.3. Các module chính", 1)
para(doc, "Hệ thống gồm các module chính: module tra cứu vé (Flights, Trains, Buses), module so sánh giá và xu hướng (Prices), module gợi ý lộ trình (RouteOptimizer), module đặt vé và thanh toán (Bookings, Payments), module xác thực (Auth), module cảnh báo giá (PriceAlerts), module quản trị (Admin), và các module hỗ trợ khác (Reviews, PromoCodes, Notifications, Subscriptions...).")

heading(doc, "3.4. Các chức năng chính", 1)

heading(doc, "3.4.1. Tra cứu chuyến bay", 2)
para(doc, "Mục đích: tìm kiếm chuyến bay theo tuyến đường, ngày đi với các bộ lọc và sắp xếp.", bold=False)
para(doc, "Đầu vào: điểm đi, điểm đến, ngày đi, các bộ lọc (giá, hãng bay, giờ, hạng ghế, số ghế), sắp xếp, phân trang.")
para(doc, "Xử lý: backend truy vấn bảng Flights theo tuyến và ngày, áp dụng bộ lọc, sắp xếp và phân trang, trả về danh sách.")
para(doc, "Kết quả: danh sách chuyến bay phù hợp kèm giá và thông tin hãng bay.")
para(doc, "File liên quan: FlightsController, PricesController, trang SearchFlights, component FlightCard.")
figure_placeholder(doc, "[CHÈN ẢNH: Giao diện tìm kiếm chuyến bay]")
figure_caption(doc, "Hình 3.1. Giao diện tìm kiếm chuyến bay")

heading(doc, "3.4.2. So sánh giá và biểu đồ xu hướng", 2)
para(doc, "Mục đích: so sánh giá vé giữa máy bay, tàu hỏa, xe khách trên cùng tuyến và hiển thị xu hướng giá.")
para(doc, "Đầu vào: điểm đi, điểm đến, ngày đi; loại lộ trình một chiều hoặc khứ hồi.")
para(doc, "Xử lý: backend gộp dữ liệu từ ba bảng Flights, Trains, Buses, sắp xếp theo giá, tính giá thấp nhất/trung bình/cao nhất; riêng gói miễn phí chỉ so sánh một hãng bay rẻ nhất. Dữ liệu xu hướng giá được truy vấn từ bảng PriceHistories.")
para(doc, "Kết quả: kết quả so sánh ba phương tiện và biểu đồ xu hướng giá (recharts).")
para(doc, "File liên quan: PricesController, PriceHistoryService, trang PriceComparison, hook usePriceStream.")
figure_placeholder(doc, "[CHÈN ẢNH: Giao diện so sánh giá và biểu đồ xu hướng]")
figure_caption(doc, "Hình 3.2. Giao diện so sánh giá và biểu đồ xu hướng")

heading(doc, "3.4.3. Gợi ý lộ trình kết hợp tối ưu", 2)
para(doc, "Mục đích: gợi ý lộ trình kết hợp nhiều phương tiện tối ưu chi phí, ví dụ đi máy bay về tàu hỏa.")
para(doc, "Đầu vào: điểm đi, điểm đến, ngày đi, ngày về, tiêu chí tối ưu (rẻ nhất, nhanh nhất, ít trung chuyển, đến sớm nhất).")
para(doc, "Xử lý: RouteOptimizerService sinh các lộ trình một chặng, hai chặng (qua một trung tâm) và ba chặng (qua hai trung tâm), kết hợp ba loại phương tiện, kiểm tra thời gian trung chuyển hợp lệ, tính tổng chi phí và thời gian, sắp xếp theo tiêu chí.")
para(doc, "Kết quả: danh sách lộ trình tối ưu kèm chi tiết từng chặng.")
para(doc, "File liên quan: RouteOptimizerService, PricesController, trang OptimalRoute.")
figure_placeholder(doc, "[CHÈN ẢNH: Giao diện gợi ý lộ trình tối ưu]")
figure_caption(doc, "Hình 3.3. Giao diện gợi ý lộ trình tối ưu")

heading(doc, "3.4.4. Đặt vé và chuyển hướng đặt vé chính thức", 2)
para(doc, "Mục đích: cho phép người dùng đặt vé trực tiếp tại hệ thống hoặc chuyển hướng đến trang đặt vé chính thức của hãng.")
para(doc, "Đầu vào: chuyến vé đã chọn, thông tin hành khách, số lượng.")
para(doc, "Xử lý: hệ thống tạo đặt chỗ (Booking). Khi người dùng chọn \u201cĐặt tại website chính thức\u201d, hệ thống mở trang web của hãng (Vietnam Airlines, VietJet, Đường sắt Việt Nam, Phương Trang...) kèm deep link; khi chọn \u201cĐặt tại Vé247\u201d, hệ thống thực hiện đặt chỗ trực tiếp.")
para(doc, "Kết quả: đặt chỗ thành công hoặc được chuyển hướng đến trang đặt vé chính thức.")
para(doc, "File liên quan: BookingsController, BookingOptionsModal, trang BookingPage, PaymentPage.")
figure_placeholder(doc, "[CHÈN ẢNH: Giao diện chọn nơi đặt vé]")
figure_caption(doc, "Hình 3.4. Giao diện chọn nơi đặt vé")

heading(doc, "3.4.5. Thanh toán", 2)
para(doc, "Mục đích: thanh toán đặt chỗ qua nhiều cổng thanh toán.")
para(doc, "Đầu vào: đặt chỗ, phương thức thanh toán (VNPay, PayOS, MoMo, ZaloPay, VietQR, thẻ, chuyển khoản).")
para(doc, "Xử lý: backend tạo yêu cầu thanh toán theo cổng, người dùng được dẫn đến cổng thanh toán, hệ thống xác minh kết quả qua trang return/IPN và cập nhật trạng thái đặt chỗ.")
para(doc, "Kết quả: cập nhật trạng thái thanh toán và đặt chỗ.")
para(doc, "File liên quan: PaymentsController, VnPayService, PayOSService, MoMoService, ZaloPayService, VietQrService, trang PaymentPage.")
figure_placeholder(doc, "[CHÈN ẢNH: Giao diện thanh toán]")
figure_caption(doc, "Hình 3.5. Giao diện thanh toán")

heading(doc, "3.4.6. Cảnh báo giá", 2)
para(doc, "Mục đích: cho phép người dùng đặt cảnh báo giá cho một tuyến và nhận thông báo khi giá giảm.")
para(doc, "Đầu vào: email, tuyến đường, mức giá mục tiêu.")
para(doc, "Xử lý: hệ thống lưu cảnh báo, định kỳ kiểm tra giá so với mục tiêu, tạo thông báo và gửi email khi giá đạt điều kiện.")
para(doc, "Kết quả: người dùng nhận thông báo và email khi giá giảm.")
para(doc, "File liên quan: PriceAlertController, PriceAlertService, EmailService.")
figure_placeholder(doc, "[CHÈN ẢNH: Giao diện cảnh báo giá]")
figure_caption(doc, "Hình 3.6. Giao diện đặt cảnh báo giá")

heading(doc, "3.4.7. Quản trị hệ thống", 2)
para(doc, "Mục đích: quản lý dữ liệu và vận hành hệ thống.")
para(doc, "Đầu vào: đăng nhập quản trị viên.")
para(doc, "Xử lý: quản lý chuyến bay/tàu/xe (thêm, sửa, xóa, nhập/xuất), quản lý người dùng, đặt chỗ, mã giảm giá, gói VIP, thống kê, phát thông báo.")
para(doc, "Kết quả: dữ liệu hệ thống được quản lý và thống kê.")
para(doc, "File liên quan: AdminController, các trang trong thư mục admin/pages, AdminDashboard.")
figure_placeholder(doc, "[CHÈN ẢNH: Giao diện quản trị]")
figure_caption(doc, "Hình 3.7. Giao diện quản trị hệ thống")

heading(doc, "3.5. Thuật toán và xử lý tiêu biểu", 1)
para(doc, "Một số xử lý tiêu biểu trong hệ thống gồm: tìm kiếm và lọc dữ liệu chuyến bay/tàu/xe với phân trang; sắp xếp theo giá, thời gian khởi hành, thời lượng; thuật toán gợi ý lộ trình tối ưu kết hợp nhiều phương tiện (RouteOptimizer) sinh các lộ trình một, hai, ba chặng và tối ưu theo tiêu chí; xử lý thanh toán đa cổng với xác minh chữ ký; và cập nhật giá thời gian thực mô phỏng bằng SignalR.")

heading(doc, "3.6. API", 1)
table_caption(doc, "Bảng 3.2. Danh sách API chính")
make_table(doc,
    ["STT", "API", "Method", "Endpoint", "Mục đích"],
    [["1", "Chuyến bay", "GET", "/api/flights", "Tra cứu chuyến bay"],
     ["2", "Tàu hỏa", "GET", "/api/trains", "Tra cứu tàu hỏa"],
     ["3", "Xe khách", "GET", "/api/buses", "Tra cứu xe khách"],
     ["4", "So sánh giá", "GET", "/api/prices/compare", "So sánh giá các phương tiện"],
     ["5", "Xu hướng giá", "GET", "/api/prices/trends", "Xu hướng giá"],
     ["6", "Dự đoán giá", "GET", "/api/prices/predict", "Dự đoán giá"],
     ["7", "Lộ trình tối ưu", "POST", "/api/prices/optimal-route", "Gợi ý lộ trình tối ưu"],
     ["8", "Đặt chỗ", "POST", "/api/bookings", "Tạo đặt chỗ"],
     ["9", "Thanh toán", "POST", "/api/bookings/{id}/pay", "Thanh toán đặt chỗ"],
     ["10", "Đăng nhập", "POST", "/api/auth/login", "Đăng nhập"],
     ["11", "Đăng ký", "POST", "/api/auth/register", "Đăng ký"],
     ["12", "Cảnh báo giá", "POST", "/api/price-alerts", "Tạo cảnh báo giá"],
     ["13", "Quản trị", "GET", "/api/admin/dashboard", "Dữ liệu dashboard"]],
    widths=[1.5, 3.5, 1.8, 5.2, 4.0])

heading(doc, "3.7. Authentication và Authorization", 1)
para(doc, "Hệ thống hỗ trợ đăng ký và đăng nhập người dùng. Mật khẩu được băm bằng BCrypt trước khi lưu. Người dùng có thể đăng nhập bằng email/mật khẩu hoặc qua OAuth (Clerk). Quyền quản trị được phân biệt qua vai trò (role) của người dùng; các API quản trị yêu cầu người dùng có vai trò Admin, và giao diện quản trị được bảo vệ bởi AdminGuard phía frontend.")

heading(doc, "3.8. Kiểm thử", 1)
para(doc, "Dự án có các bài kiểm thử đơn vị sử dụng xUnit cho các service như chatbot (ChatBotServiceTests), xử lý thanh toán PayOS (PayOSReturnTests) và gợi ý lộ trình (RouteOptimizerServiceTests). Bên cạnh đó, các chức năng được kiểm thử thủ công thông qua việc chạy hệ thống và kiểm tra API.")
table_caption(doc, "Bảng 3.3. Danh sách các bài kiểm thử đơn vị")
make_table(doc,
    ["STT", "File kiểm thử", "Đối tượng", "Mục đích"],
    [["1", "ChatBotServiceTests.cs", "ChatBotService", "Kiểm thử logic chatbot rule-based"],
     ["2", "PayOSReturnTests.cs", "PayOSService", "Kiểm thử xử lý kết quả thanh toán PayOS"],
     ["3", "RouteOptimizerServiceTests.cs", "RouteOptimizerService", "Kiểm thử logic gợi ý lộ trình"]],
    widths=[1.5, 5.5, 4.0, 4.0])

heading(doc, "3.9. Triển khai hệ thống", 1)
heading(doc, "3.9.1. Định hướng triển khai trên Vercel", 2)
para(doc, "Vercel là nền tảng triển khai frontend phổ biến, phù hợp với ứng dụng React/Vite. Frontend dùng Vite với lệnh build là npm run build, thư mục đầu ra là dist, phù hợp để triển khai trên Vercel. Các biến môi trường cần cấu hình như VITE_API_URL và VITE_WS_URL. Tuy nhiên, backend sử dụng ASP.NET Core và cơ sở dữ liệu SQL Server, không tương thích trực tiếp với nền tảng serverless Vercel, nên backend cần được triển khai riêng (ví dụ trên Azure).")
para(doc, "Tại thời điểm hoàn thiện báo cáo, hệ thống chưa được triển khai thực tế thành công trên Vercel. Nội dung này trình bày phương án và định hướng triển khai.")

heading(doc, "3.10. Giao diện chương trình", 1)
para(doc, "Hệ thống có các giao diện chính gồm trang chủ, tìm kiếm chuyến bay, tìm kiếm tàu hỏa, tìm kiếm xe khách, so sánh giá, gợi ý lộ trình, đặt vé, thanh toán, đăng nhập/đăng ký, quản lý đặt chỗ, hồ sơ cá nhân, và trang quản trị. Dưới đây là danh sách các giao diện chính cần chụp ảnh bổ sung.")

heading(doc, "3.10.1. Trang chủ", 2)
para(doc, "Mục đích: giới thiệu và cho phép tìm kiếm nhanh. Các thành phần chính: thanh tìm kiếm, banner, gợi ý tuyến phổ biến. Chức năng: tìm kiếm vé, xem ưu đãi.")
figure_placeholder(doc, "[CHÈN ẢNH: Giao diện trang chủ]")
figure_caption(doc, "Hình 3.8. Giao diện trang chủ")

heading(doc, "3.10.2. Trang đăng nhập/đăng ký", 2)
para(doc, "Mục đích: xác thực người dùng. Các thành phần chính: biểu mẫu đăng nhập, đăng ký, đăng nhập qua OAuth.")
figure_placeholder(doc, "[CHÈN ẢNH: Giao diện đăng nhập/đăng ký]")
figure_caption(doc, "Hình 3.9. Giao diện đăng nhập/đăng ký")

heading(doc, "3.10.3. Trang tìm kiếm tàu hỏa", 2)
para(doc, "Mục đích: tra cứu chuyến tàu. Các thành phần chính: biểu mẫu tìm kiếm, danh sách kết quả, bộ lọc.")
figure_placeholder(doc, "[CHÈN ẢNH: Giao diện tìm kiếm tàu hỏa]")
figure_caption(doc, "Hình 3.10. Giao diện tìm kiếm tàu hỏa")

heading(doc, "3.10.4. Trang tìm kiếm xe khách", 2)
para(doc, "Mục đích: tra cứu chuyến xe khách. Các thành phần chính: biểu mẫu tìm kiếm, danh sách kết quả, bộ lọc.")
figure_placeholder(doc, "[CHÈN ẢNH: Giao diện tìm kiếm xe khách]")
figure_caption(doc, "Hình 3.11. Giao diện tìm kiếm xe khách")

heading(doc, "3.10.5. Trang đặt vé", 2)
para(doc, "Mục đích: nhập thông tin và tạo đặt chỗ. Các thành phần chính: thông tin chuyến, thông tin hành khách, thông tin liên hệ.")
figure_placeholder(doc, "[CHÈN ẢNH: Giao diện đặt vé]")
figure_caption(doc, "Hình 3.12. Giao diện đặt vé")

heading(doc, "3.10.6. Trang quản lý đặt chỗ", 2)
para(doc, "Mục đích: xem và quản lý các đặt chỗ của người dùng. Các thành phần chính: danh sách đặt chỗ, chi tiết, hủy vé.")
figure_placeholder(doc, "[CHÈN ẢNH: Giao diện quản lý đặt chỗ]")
figure_caption(doc, "Hình 3.13. Giao diện quản lý đặt chỗ")

heading(doc, "3.10.7. Trang quản trị", 2)
para(doc, "Mục đích: quản lý dữ liệu hệ thống. Các thành phần chính: dashboard, danh sách chuyến bay/tàu/xe, người dùng, đặt chỗ, thống kê.")
figure_placeholder(doc, "[CHÈN ẢNH: Giao diện dashboard quản trị]")
figure_caption(doc, "Hình 3.14. Giao diện dashboard quản trị")
page_break(doc)

# ============================ PHẦN BA: KẾT LUẬN ============================
sec = doc.add_section(WD_SECTION.NEW_PAGE)
set_section(sec)
sec.footer.is_linked_to_previous = False
enable_page_number(sec)

chapter_title(doc, "PHẦN BA: KẾT LUẬN")

heading(doc, "1. Kết quả đạt được", 1)
heading(doc, "1.1. Kết quả về chức năng", 2)
para(doc, "Hệ thống đã xây dựng và vận hành được các chức năng cốt lõi: tra cứu chuyến bay, tàu hỏa, xe khách với bộ lọc và sắp xếp; so sánh giá và hiển thị biểu đồ xu hướng giá; gợi ý lộ trình kết hợp tối ưu chi phí; đặt vé trực tiếp và chuyển hướng đến trang đặt vé chính thức; thanh toán đa cổng; cảnh báo giá; xác thực người dùng; và trang quản trị.")
heading(doc, "1.2. Kết quả về công nghệ", 2)
para(doc, "Đồ án đã ứng dụng thành công bộ công nghệ web hiện đại gồm React, Vite, Tailwind CSS, ASP.NET Core, Entity Framework Core, SignalR và SQL Server, thể hiện khả năng xây dựng ứng dụng full-stack.")
heading(doc, "1.3. Kết quả về giao diện", 2)
para(doc, "Giao diện được thiết kế hiện đại, thân thiện, responsive và nhất quán giữa các trang, hỗ trợ chế độ sáng/tối.")
heading(doc, "1.4. Kết quả về kiến thức và kinh nghiệm", 2)
para(doc, "Qua đồ án, sinh viên đã củng cố kiến thức về phân tích thiết kế hệ thống, lập trình full-stack, quản lý cơ sở dữ liệu, tích hợp thanh toán và kiểm thử phần mềm.")

heading(doc, "2. Hạn chế", 1)
para(doc, "Một số hạn chế của hệ thống: dữ liệu giá vé hiện tại là dữ liệu mẫu được khởi tạo vào cơ sở dữ liệu, chưa tích hợp trực tiếp với API giá vé thực tế do chi phí; tính năng \u201cthời gian thực\u201d hiện được mô phỏng bằng cách tự làm thay đổi giá định kỳ qua SignalR chứ không phải dữ liệu cập nhật từ hãng; một số cổng thanh toán (MoMo, ZaloPay, VietQR) chưa có thông tin cấu hình đầy đủ; hệ thống chưa được triển khai lên môi trường production; phạm vi kiểm thử tự động còn hạn chế.")

heading(doc, "3. Hướng phát triển", 1)
para(doc, "Hướng phát triển của đề tài gồm: hoàn thiện việc tích hợp dữ liệu từ nhiều nguồn thực tế thông qua lớp nguồn dữ liệu (provider) gồm nguồn cơ sở dữ liệu, nguồn giá tàu từ website Đường sắt Việt Nam và nguồn dữ liệu tham chiếu sân bay, hãng bay; hoàn thiện triển khai trên Vercel và Azure; bổ sung thêm kiểm thử tự động và CI/CD; tối ưu hiệu năng và bảo mật; phát triển ứng dụng di động.")
para(doc, "Trong đó, việc tích hợp nhiều nguồn dữ liệu (đã được định hướng) là hướng phát triển chính để hệ thống đạt được dữ liệu giá vé gần với thực tế hơn.")

para(doc, "")
chapter_title(doc, "GHI CHÚ CÁC THUẬT NGỮ")
table_caption(doc, "Bảng 4.1. Danh mục thuật ngữ viết tắt")
make_table(doc,
    ["STT", "Thuật ngữ", "Viết đầy đủ", "Giải thích"],
    [["1", "API", "Application Programming Interface", "Giao diện lập trình ứng dụng"],
     ["2", "UI", "User Interface", "Giao diện người dùng"],
     ["3", "UX", "User Experience", "Trải nghiệm người dùng"],
     ["4", "CRUD", "Create, Read, Update, Delete", "Thao tác cơ bản trên dữ liệu"],
     ["5", "ORM", "Object-Relational Mapping", "Ánh xạ đối tượng - quan hệ"],
     ["6", "SQL", "Structured Query Language", "Ngôn ngữ truy vấn cơ sở dữ liệu"],
     ["7", "REST", "Representational State Transfer", "Kiến trúc API web"],
     ["8", "DB", "Database", "Cơ sở dữ liệu"]],
    widths=[1.2, 2.5, 6.0, 5.3])
page_break(doc)

# ============================ PHỤ LỤC ============================
sec = doc.add_section(WD_SECTION.NEW_PAGE)
set_section(sec)
sec.footer.is_linked_to_previous = False
enable_page_number(sec)

chapter_title(doc, "PHỤ LỤC")

chapter_title(doc, "PHỤ LỤC A: HƯỚNG DẪN CÀI ĐẶT DEMO")
heading(doc, "A.1. Yêu cầu môi trường", 1)
para(doc, "Hệ thống yêu cầu: .NET SDK 10 trở lên cho backend; Node.js và npm cho frontend; kết nối đến SQL Server (Azure SQL Server) với chuỗi kết nối hợp lệ.")
heading(doc, "A.2. Cài đặt dependency", 1)
para(doc, "Frontend: chạy lệnh npm install trong thư mục frontend. Backend: chạy dotnet restore trong thư mục backend/FlightAggregatorApi.")
heading(doc, "A.3. Cấu hình biến môi trường", 1)
para(doc, "Frontend cấu hình biến VITE_API_URL và VITE_WS_URL trong file frontend/.env trỏ đến backend. Backend cấu hình chuỗi kết nối AzureSqlDb trong appsettings.json.")
heading(doc, "A.4. Cấu hình database", 1)
para(doc, "Backend tự động tạo bảng và seed dữ liệu mẫu khi khởi động thông qua DatabaseInitializerService. Cần đảm bảo chuỗi kết nối đến SQL Server hợp lệ.")
heading(doc, "A.5. Chạy backend", 1)
para(doc, "Chạy backend bằng lệnh dotnet run --project backend/FlightAggregatorApi hoặc dùng script scripts/start-backend.bat. Backend chạy trên cổng 5000.")
heading(doc, "A.6. Chạy frontend", 1)
para(doc, "Chạy frontend bằng lệnh npm run dev trong thư mục frontend. Frontend chạy trên cổng 5173.")
heading(doc, "A.7. Chạy ứng dụng", 1)
para(doc, "Mở trình duyệt truy cập http://localhost:5173 để sử dụng hệ thống.")
heading(doc, "A.8. Kiểm tra kết nối API", 1)
para(doc, "Kiểm tra backend qua endpoint http://localhost:5000/health, trả về trạng thái healthy nếu hệ thống hoạt động bình thường.")
heading(doc, "A.9. Các lỗi thường gặp", 1)
para(doc, "Một số lỗi thường gặp: cổng 5000 đang bị chiếm giữ do instance backend cũ (dùng script stop-backend.bat hoặc tự động trong run-backend.bat); lỗi kết nối cơ sở dữ liệu do chuỗi kết nối sai; thiếu biến môi trường VITE_API_URL.")

heading(doc, "Tài khoản kiểm thử", 1)
table_caption(doc, "Bảng 4.2. Tài khoản kiểm thử")
make_table(doc,
    ["Loại tài khoản", "Username/Email", "Password", "Quyền"],
    [["Người dùng demo", "user@example.com", "123456", "User"],
     ["Quản trị viên demo", "admin@ve247.vn", "Admin123", "Admin"]],
    widths=[4.0, 5.5, 3.0, 2.5])
page_break(doc)

chapter_title(doc, "PHỤ LỤC B: HƯỚNG DẪN SỬ DỤNG")
para(doc, "Dưới đây là hướng dẫn sử dụng một số chức năng tiêu biểu của hệ thống.")
para(doc, "Bước 1: Truy cập http://localhost:5173 để mở trang chủ.")
para(doc, "Bước 2: Tại thanh tìm kiếm, chọn loại phương tiện, nhập điểm đi, điểm đến, ngày đi rồi bấm tìm kiếm.")
para(doc, "Bước 3: Chọn một kết quả vé, xem chi tiết và bấm Đặt vé.")
para(doc, "Bước 4: Trong màn hình đặt vé, nhập thông tin hành khách và chọn nơi đặt (đặt tại Vé247 hoặc chuyển hướng đến trang chính thức).")
para(doc, "Bước 5: Thanh toán đặt chỗ theo phương thức đã chọn và kiểm tra kết quả.")
para(doc, "Ngoài ra, người dùng có thể dùng chức năng so sánh giá và gợi ý lộ trình trên các trang tương ứng để chọn phương án tối ưu.")

doc.add_page_break()

# ============================ TÀI LIỆU THAM KHẢO ============================
sec = doc.add_section(WD_SECTION.NEW_PAGE)
set_section(sec)
sec.footer.is_linked_to_previous = False
enable_page_number(sec)

chapter_title(doc, "TÀI LIỆU THAM KHẢO")
heading(doc, "1. Sách, giáo trình", 1)
heading(doc, "1.1. Tiếng Việt", 2)
para(doc, "Các giáo trình về phân tích thiết kế hệ thống thông tin, cơ sở dữ liệu và phát triển ứng dụng web được sử dụng trong chương trình đào tạo.")
heading(doc, "1.2. Tiếng Anh", 2)
para(doc, "Tài liệu về React, ASP.NET Core và Entity Framework Core được tham khảo từ tài liệu chính thức của các nền tảng.")
heading(doc, "2. Các website tham khảo", 1)
para(doc, "[1] React Documentation; Online, truy cập ngày 17/08/2026, tại https://react.dev")
para(doc, "[2] Vite Documentation; Online, truy cập ngày 17/08/2026, tại https://vite.dev")
para(doc, "[3] ASP.NET Core Documentation; Online, truy cập ngày 17/08/2026, tại https://learn.microsoft.com/aspnet/core")
para(doc, "[4] Entity Framework Core Documentation; Online, truy cập ngày 17/08/2026, tại https://learn.microsoft.com/ef/core")
para(doc, "[5] Microsoft SQL Server Documentation; Online, truy cập ngày 17/08/2026, tại https://learn.microsoft.com/sql")
para(doc, "[6] Tailwind CSS Documentation; Online, truy cập ngày 17/08/2026, tại https://tailwindcss.com")
para(doc, "[7] recharts Documentation; Online, truy cập ngày 17/08/2026, tại https://recharts.org")

doc.save(OUT)
print("SAVED:", OUT)